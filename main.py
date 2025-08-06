print("Starting backend...")
# FastAPI backend code goes here (will provide full code in next steps)
from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks, HTTPException, Depends, Query
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import os
from typing import List, Optional, Dict
import pdfplumber
import docx
from PIL import Image
import pytesseract
from dotenv import load_dotenv
import google.generativeai as genai
from openai import AsyncOpenAI
import asyncio
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import uuid
from datetime import datetime
import hashlib
from docx2pdf import convert
from pydantic import BaseModel, EmailStr
import requests

app = FastAPI()

# Allow CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
REPORT_DIR = "reports"
USERS_DIR = "users"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)
os.makedirs(USERS_DIR, exist_ok=True)

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)
gemini_model = genai.GenerativeModel('gemini-2.5-flash')
llama_client = AsyncOpenAI(
    api_key=OPENAI_API_KEY,
    base_url="https://api.together.xyz/v1"  # IMPORTANT: Change if not using Together.ai
)

# User management
users_file = os.path.join(USERS_DIR, "users.json")
user_data_file = os.path.join(USERS_DIR, "user_data.json")

def load_users():
    if os.path.exists(users_file):
        with open(users_file, 'r') as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(users_file, 'w') as f:
        json.dump(users, f, indent=2)

def load_user_data():
    if os.path.exists(user_data_file):
        with open(user_data_file, 'r') as f:
            return json.load(f)
    return {}

def save_user_data(user_data):
    with open(user_data_file, 'w') as f:
        json.dump(user_data, f, indent=2)

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed

def generate_token() -> str:
    return str(uuid.uuid4())

def get_user_from_token(token: str) -> Optional[Dict]:
    users = load_users()
    for user_id, user in users.items():
        if user.get('token') == token:
            return {'id': user_id, **user}
    return None

# Authentication dependency
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer())):
    token = credentials.credentials
    user = get_user_from_token(token)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid token")
    return user

@app.post("/auth/signup")
async def signup(email: str = Form(...), password: str = Form(...), name: str = Form(...)):
    users = load_users()
    
    # Check if user already exists
    for user in users.values():
        if user['email'] == email:
            raise HTTPException(status_code=400, detail="User already exists")
    
    # Create new user
    user_id = str(uuid.uuid4())
    token = generate_token()
    
    new_user = {
        'email': email,
        'password': hash_password(password),
        'name': name,
        'token': token,
        'created_at': datetime.now().isoformat()
    }
    
    users[user_id] = new_user
    save_users(users)
    
    # Initialize user data
    user_data = load_user_data()
    user_data[user_id] = {
        'documents_analyzed': 0,
        'reports_generated': 0,
        'last_analysis': None,
        'analysis_history': []
    }
    save_user_data(user_data)
    
    return {
        'token': token,
        'user': {
            'id': user_id,
            'email': email,
            'name': name
        }
    }

@app.post("/auth/login")
async def login(email: str = Form(...), password: str = Form(...)):
    users = load_users()
    
    for user_id, user in users.items():
        if user['email'] == email and verify_password(password, user['password']):
            # Generate new token
            token = generate_token()
            user['token'] = token
            save_users(users)
            
            return {
                'token': token,
                'user': {
                    'id': user_id,
                    'email': user['email'],
                    'name': user['name']
                }
            }
    
    raise HTTPException(status_code=401, detail="Invalid credentials")

@app.get("/user/profile")
async def get_user_profile(current_user: Dict = Depends(get_current_user)):
    return {
        'id': current_user['id'],
        'email': current_user['email'],
        'name': current_user['name'],
        'created_at': current_user['created_at']
    }

@app.get("/user/stats")
async def get_user_stats(current_user: Dict = Depends(get_current_user)):
    user_data = load_user_data()
    user_stats = user_data.get(current_user['id'], {
        'documents_analyzed': 0,
        'reports_generated': 0,
        'last_analysis': None,
        'analysis_history': []
    })
    
    return user_stats

@app.get("/user/history")
async def get_user_history(current_user: Dict = Depends(get_current_user)):
    user_data = load_user_data()
    user_stats = user_data.get(current_user['id'], {})
    return user_stats.get('analysis_history', [])

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    return pytesseract.image_to_string(image)

@app.post("/upload")
async def upload_files(
    document: UploadFile = File(...),
    screenshots: Optional[List[UploadFile]] = File(None),
    chapter: Optional[str] = Form(None),
    current_user: Dict = Depends(get_current_user)
):
    # Create user-specific upload directory
    user_upload_dir = os.path.join(UPLOAD_DIR, current_user['id'])
    os.makedirs(user_upload_dir, exist_ok=True)
    
    # Save document
    doc_path = os.path.join(user_upload_dir, document.filename)
    with open(doc_path, "wb") as f:
        f.write(await document.read())
    
    screenshot_paths = []
    if screenshots:
        for shot in screenshots:
            shot_path = os.path.join(user_upload_dir, shot.filename)
            with open(shot_path, "wb") as f:
                f.write(await shot.read())
            screenshot_paths.append(shot_path)
    
    # Return file info and analysis trigger token (placeholder)
    return {"document": doc_path, "screenshots": screenshot_paths, "chapter": chapter, "token": "demo-token"}

def setup_document_styles(document):
    """Sets up custom styles for the document."""
    styles = document.styles
    try:
        # Base style
        base_style = styles['Normal']
        font = base_style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title style
        title_style = styles.add_style('ReportTitle', WD_STYLE_TYPE.PARAGRAPH)
        font = title_style.font
        font.name = 'Calibri'
        font.size = Pt(28)
        font.bold = True
        font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        # Heading 1 style
        h1_style = styles.add_style('ReportHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = styles['Heading 1']
        font = h1_style.font
        font.name = 'Calibri'
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        p_format = h1_style.paragraph_format
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(6)
    except Exception as e:
        print(f"An error occurred during style setup: {e}")

@app.post("/analyze")
async def analyze(
    token: str = Form(...),
    current_user: Dict = Depends(get_current_user)
):
    user_id = current_user['id']
    user_upload_dir = os.path.join(UPLOAD_DIR, user_id)
    
    # This is a simplified lookup based on the most recent files.
    # A more robust system would use a database to track file uploads per user.
    all_user_files = [os.path.join(user_upload_dir, f) for f in os.listdir(user_upload_dir)]
    if not all_user_files:
        raise HTTPException(status_code=404, detail="No documents found for this user to analyze.")

    all_user_files.sort(key=os.path.getmtime, reverse=True)
    
    doc_path = next((f for f in all_user_files if f.lower().endswith(('.pdf', '.docx'))), None)
    screenshot_paths = [f for f in all_user_files if f.lower().endswith(('.png', '.jpg', '.jpeg'))][:5] # Limit to 5 screenshots

    if not doc_path:
        raise HTTPException(status_code=404, detail="A compatible document (.pdf, .docx) for analysis not found in recent uploads.")

    original_filename = os.path.basename(doc_path)
    file_type = doc_path.split('.')[-1].lower()
    
    # 1. Extract text
    doc_text = ""
    if file_type == 'pdf':
        doc_text = extract_text_from_pdf(doc_path)
    elif file_type == 'docx':
        doc_text = extract_text_from_docx(doc_path)

    screenshot_texts = [extract_text_from_image(p) for p in screenshot_paths]
    
    # 2. Perform AI analysis concurrently
    analysis_tasks = {
        "summary": asyncio.to_thread(llama3_summarize, doc_text),
        "grammar": asyncio.to_thread(llama3_grammar_correct, doc_text),
        "suggestions": asyncio.to_thread(llama3_suggestions, doc_text),
        "inconsistencies": asyncio.to_thread(llama3_inconsistencies, doc_text, screenshot_texts),
        "repetition": asyncio.to_thread(llama3_check_for_repetition, doc_text),
        "internal_inconsistencies": asyncio.to_thread(llama3_check_internal_inconsistencies, doc_text)
    }
    
    results = await asyncio.gather(*analysis_tasks.values())
    summary, grammar_correction, suggestions, inconsistencies, repetition_check, internal_inconsistencies = results

    # 3. Generate DOCX report in a user-specific directory
    user_report_dir = os.path.join(REPORT_DIR, user_id)
    os.makedirs(user_report_dir, exist_ok=True)
    
    doc_id = str(uuid.uuid4()).split('-')[0]
    report_filename = f"report_{user_id}_{doc_id}.docx"
    report_path = os.path.join(user_report_dir, report_filename)
    
    document = Document()
    setup_document_styles(document)
    
    document.add_heading('AI Analysis Report', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Analysis for: {original_filename}", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Analyzed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_section(title, content):
        if content and content.strip():
            document.add_heading(title, level=1)
            lines = content.strip().split('\n')
            for line in lines:
                stripped = line.strip()
                # Bullet points ("* " or "- ")
                if stripped.startswith("* ") or stripped.startswith("- "):
                    document.add_paragraph(stripped[2:], style='List Bullet')
                # Numbered list (e.g., "1. ", "2. ")
                elif len(stripped) > 2 and stripped[:2].isdigit() and stripped[2:4] == ". ":
                    document.add_paragraph(stripped[4:], style='List Number')
                # Markdown-style subheading (e.g., "### Heading")
                elif stripped.startswith("### "):
                    document.add_heading(stripped[4:], level=2)
                # Normal paragraph
                else:
                    document.add_paragraph(stripped)

    add_section("Summary", summary)
    add_section("Grammar Corrections", grammar_correction)
    add_section("Improvement Suggestions", suggestions)
    add_section("Screenshot Inconsistencies", inconsistencies)
    add_section("Repetitive Content Check", repetition_check)
    add_section("Internal Inconsistencies Check", internal_inconsistencies)
    
    document.save(report_path)

    # 4. Store analysis results in user's history
    user_data = load_user_data()
    report_id = f"{user_id}_{doc_id}"
    
    analysis_entry = {
        'id': report_id,
        'original_filename': original_filename,
        'timestamp': datetime.now().isoformat(),
        'summary': summary,
        'grammar_correction': grammar_correction,
        'suggestions': suggestions,
        'inconsistencies': inconsistencies,
        'repetition_check': repetition_check,
        'internal_inconsistencies': internal_inconsistencies,
        'report_path': report_path,
        'file_type': file_type,
    }
    
    if user_id in user_data:
        user_data[user_id]['documents_analyzed'] += 1
        user_data[user_id]['reports_generated'] += 1
        user_data[user_id]['last_analysis'] = datetime.now().isoformat()
        user_data[user_id]['analysis_history'].insert(0, analysis_entry)
        user_data[user_id]['analysis_history'] = user_data[user_id]['analysis_history'][:20]
    
    save_user_data(user_data)
    
    return {
        "message": "Analysis complete",
        "report_id": report_id,
        "results": analysis_entry
    }

@app.delete("/analysis/{report_id}")
async def delete_analysis(report_id: str, current_user: Dict = Depends(get_current_user)):
    user_id = current_user['id']
    
    # Ensure user can only delete their own reports
    if not report_id.startswith(user_id):
        raise HTTPException(status_code=403, detail="Access denied: You can only delete your own analysis history.")

    user_data = load_user_data()
    user_history = user_data.get(user_id, {}).get('analysis_history', [])
    
    analysis_to_delete = None
    for item in user_history:
        if item['id'] == report_id:
            analysis_to_delete = item
            break
            
    if not analysis_to_delete:
        raise HTTPException(status_code=404, detail="Analysis not found in history.")
        
    # Remove from history
    user_history.remove(analysis_to_delete)
    user_data[user_id]['analysis_history'] = user_history
    save_user_data(user_data)
    
    # Optionally, delete the physical report file from user's directory
    user_report_dir = os.path.join(REPORT_DIR, user_id)
    report_path_docx = os.path.join(user_report_dir, f"report_{report_id}.docx")
    report_path_pdf = os.path.join(user_report_dir, f"report_{report_id}.pdf")
    
    if os.path.exists(report_path_docx):
        os.remove(report_path_docx)
    if os.path.exists(report_path_pdf):
        os.remove(report_path_pdf)
        
    return {"message": "Analysis deleted successfully"}

@app.get("/analysis/{report_id}")
async def get_analysis_details(report_id: str, current_user: Dict = Depends(get_current_user)):
    # Ensure user can only access their own reports
    if not report_id.startswith(current_user['id']):
        raise HTTPException(status_code=403, detail="Access denied")

    user_data = load_user_data()
    user_history = user_data.get(current_user['id'], {}).get('analysis_history', [])
    
    analysis = next((item for item in user_history if item['id'] == report_id), None)

    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
        
    return analysis

@app.get("/report/{report_id}")
def get_report(
    report_id: str, 
    format: str = Query("docx", enum=["docx", "pdf"]),
    current_user: Dict = Depends(get_current_user)
):
    # Ensure user can only access their own reports
    if not report_id.startswith(current_user['id']):
        print(f"Access denied: User {current_user['id']} tried to access report {report_id}")
        raise HTTPException(status_code=403, detail="Access denied - You can only access your own reports")
    
    user_id = current_user['id']
    user_report_dir = os.path.join(REPORT_DIR, user_id)
    
    report_path_docx = os.path.join(user_report_dir, f"report_{report_id}.docx")
    
    if not os.path.exists(report_path_docx):
        print(f"Report not found: {report_path_docx}")
        raise HTTPException(status_code=404, detail="Report not found")

    if format.lower() == 'pdf':
        report_path_pdf = os.path.join(user_report_dir, f"report_{report_id}.pdf")
        
        # Convert to PDF if it doesn't exist
        if not os.path.exists(report_path_pdf):
            try:
                print(f"Converting {report_path_docx} to PDF...")
                convert(report_path_docx, report_path_pdf)
                print("Conversion complete.")
            except Exception as e:
                print(f"Error converting to PDF: {e}")
                raise HTTPException(status_code=500, detail="Failed to convert report to PDF.")
        
        return FileResponse(
            report_path_pdf, 
            media_type="application/pdf", 
            filename=f"AI_Document_Analysis_{report_id}.pdf"
        )

    # Default to DOCX
    return FileResponse(
        report_path_docx, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        filename=f"AI_Document_Analysis_{report_id}.docx"
    )

def llama3_generate(prompt, system_prompt=None):
    # Use Gemini as primary LLM service
    try:
        if GEMINI_API_KEY and GEMINI_API_KEY != "your_gemini_api_key_here":
            response = gemini_model.generate_content(prompt)
            return response.text
        else:
            return "Error: Gemini API key not configured. Please set GEMINI_API_KEY in your environment variables."
    except Exception as e:
        return f"Error using Gemini: {str(e)}"

def llama3_summarize(text):
    prompt = f"Summarize the following document:\n{text}"
    return llama3_generate(prompt)

def llama3_grammar_correct(text):
    prompt = f"Correct the grammar in the following text:\n{text}"
    return llama3_generate(prompt)

def llama3_suggestions(text):
    prompt = f"Suggest improvements for the following document:\n{text}"
    return llama3_generate(prompt)

def llama3_inconsistencies(doc_text, screenshot_texts):
    joined_screens = "\n".join(screenshot_texts)
    prompt = f"Check for inconsistencies between the following document and screenshots.\nDocument:\n{doc_text}\nScreenshots:\n{joined_screens}"
    return llama3_generate(prompt)

def llama3_check_for_repetition(text):
    prompt = f"Analyze the following text and identify any repetitive phrases, sentences, or ideas. List the redundant parts and suggest how they could be consolidated or rewritten for better clarity.\n\nText:\n{text}"
    return llama3_generate(prompt)

def llama3_check_internal_inconsistencies(text):
    prompt = f"Analyze the following document for internal inconsistencies. Check for contradictory statements, conflicting data or numbers, and inconsistencies in definitions or terminology. List any inconsistencies you find.\n\nDocument:\n{text}"
    return llama3_generate(prompt)

@app.get("/test-auth")
async def test_auth(current_user: Dict = Depends(get_current_user)):
    return {
        "message": "Authentication successful",
        "user_id": current_user['id'],
        "user_name": current_user['name']
    }

class UserUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[EmailStr] = None
    password: Optional[str] = None
    current_password: Optional[str] = None

@app.put("/user/profile")
async def update_user_profile(
    update_data: UserUpdate,
    current_user: Dict = Depends(get_current_user)
):
    users = load_users()
    user_id = current_user['id']

    if user_id not in users:
        raise HTTPException(status_code=404, detail="User not found")

    user_to_update = users[user_id]
    
    # If email or password is being changed, current password is required for verification
    if update_data.email or update_data.password:
        if not update_data.current_password:
            raise HTTPException(status_code=400, detail="Current password is required to change email or password.")
        
        if not verify_password(update_data.current_password, user_to_update['password']):
            raise HTTPException(status_code=403, detail="Incorrect current password.")

    # Update name if provided
    if update_data.name is not None and update_data.name != user_to_update['name']:
        user_to_update['name'] = update_data.name

    # Update email if provided
    if update_data.email and update_data.email != user_to_update['email']:
        for uid, user in users.items():
            if user['email'] == update_data.email and uid != user_id:
                raise HTTPException(status_code=400, detail="Email already registered by another user.")
        user_to_update['email'] = update_data.email

    # Update password if provided
    if update_data.password:
        if len(update_data.password) < 4: # Basic validation
             raise HTTPException(status_code=400, detail="Password must be at least 4 characters long.")
        user_to_update['password'] = hash_password(update_data.password)

    users[user_id] = user_to_update
    save_users(users)

    # Return the updated user data (excluding password)
    updated_user_info = user_to_update.copy()
    del updated_user_info['password']
    
    return {"message": "Profile updated successfully", "user": updated_user_info}